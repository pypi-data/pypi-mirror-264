from io import BytesIO
import fitz
import docx

class PDF:
    def __init__(self) -> None:
        self.reader = fitz
    
    def read_pdf(self, stream:bytes, zoom:int = 100) -> dict:
        doc = self.reader.Document(stream=stream, filetype="pdf")
        pages_data = {}
        for pg in range(doc.page_count):
            page = doc.load_page(pg)
            rotate = int(0)
            zoom = int(zoom)
            mat = fitz.Matrix(zoom/100.0, zoom/100.0).prerotate(rotate)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            pages_data[pg] = pix
        return pages_data

pdf = PDF()

class PDF_Convert:
    def __init__(self) -> None:
        pass
    
    @staticmethod
    def to_image(data:bytes):
        pages = pdf.read_pdf(BytesIO(data))
        return pages
    
    @staticmethod
    def to_doc(data:bytes):
        doc = pdf.reader.Document(stream=BytesIO(data), filetype="pdf")
        word = docx.Document()
        for pg in range(doc.page_count):
            page = doc.load_page(pg)  
            text = page.get_text()
            word.add_paragraph(text)
        return word
    
    @staticmethod
    def to_text(data:bytes):
        doc = pdf.reader.Document(stream=BytesIO(data), filetype="pdf")
        text = ""
        for pg in range(doc.page_count):
            page = doc.load_page(pg)  
            t = page.get_text()
            text += t
        return text
    
    @staticmethod
    def merge_pic(data:list):
        doc = pdf.reader.Document(filetype="pdf")
        for d in data:
            img = pdf.reader.Document(stream=d, filetype='JPEG')
            imgPdf = fitz.Document("pdf", img.convert_to_pdf())
            doc.insert_pdf(imgPdf)
        return doc

    
